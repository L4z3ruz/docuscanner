from flask import Flask, render_template, request, jsonify, send_from_directory, send_file
from werkzeug.exceptions import HTTPException
import cv2
import numpy as np
import pytesseract
from pytesseract.pytesseract import TesseractError, TesseractNotFoundError
import base64
import os
import io
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
SAVE_DIR = "outputs"
os.makedirs(SAVE_DIR, exist_ok=True)
API_ROUTES = {"/detect", "/scan", "/save", "/export-txt", "/export-doc"}

if os.getenv("TESSERACT_CMD"):
    pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD")
else:
    for candidate in (
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ):
        if os.path.exists(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            break

def order_points(pts):
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    return rect

def four_point_transform(image, pts):
    rect = order_points(pts)
    tl, tr, br, bl = rect
    maxWidth  = max(int(np.linalg.norm(br - bl)), int(np.linalg.norm(tr - tl)))
    maxHeight = max(int(np.linalg.norm(tr - br)), int(np.linalg.norm(tl - bl)))
    dst = np.array([[0, 0],[maxWidth-1, 0],[maxWidth-1, maxHeight-1],[0, maxHeight-1]], dtype="float32")
    M = cv2.getPerspectiveTransform(rect, dst)
    return cv2.warpPerspective(image, M, (maxWidth, maxHeight))

def detect_document_contour(frame):
    h, w = frame.shape[:2]
    img_area = h * w
    strategies = [
        (5, 30, 100, 0.02),
        (7, 50, 150, 0.02),
        (9, 10, 80,  0.03),
        (5, 80, 200, 0.02),
    ]
    for bk, clo, chi, eps in strategies:
        gray  = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        gray  = cv2.GaussianBlur(gray, (bk, bk), 0)
        edged = cv2.Canny(gray, clo, chi)
        edged = cv2.dilate(edged, None, iterations=2)
        edged = cv2.erode(edged, None, iterations=1)
        cnts, _ = cv2.findContours(edged, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        cnts = sorted(cnts, key=cv2.contourArea, reverse=True)[:10]
        for c in cnts:
            area = cv2.contourArea(c)
            if area < img_area * 0.05:
                continue
            peri   = cv2.arcLength(c, True)
            approx = cv2.approxPolyDP(c, eps * peri, True)
            if len(approx) == 4:
                return approx.reshape(4, 2).astype("float32")
    return None

def enhance(warped, mode="scan"):
    color_bytes = cv2.imencode(".png", warped)[1].tobytes()
    gray = cv2.cvtColor(warped, cv2.COLOR_BGR2GRAY)
    dil = cv2.dilate(gray, np.ones((7,7), np.uint8))
    bg  = cv2.medianBlur(dil, 21)
    sr  = cv2.absdiff(gray, bg)
    sr  = 255 - sr
    clahe    = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
    contrast = clahe.apply(sr)
    denoised = cv2.fastNlMeansDenoising(contrast, h=10)
    blurred  = cv2.GaussianBlur(denoised, (0,0), 3)
    sharp    = cv2.addWeighted(denoised, 1.5, blurred, -0.5, 0)
    up = cv2.resize(sharp, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    if mode == "scan":
        final = cv2.adaptiveThreshold(up, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10)
        kern  = np.ones((2,2), np.uint8)
        final = cv2.morphologyEx(final, cv2.MORPH_CLOSE, kern)
    else:
        final = up
    bw_bytes = cv2.imencode(".png", final)[1].tobytes()
    return color_bytes, bw_bytes, final

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/detect", methods=["POST"])
def detect():
    payload = request.get_json(silent=True) or {}
    img = _b64_to_cv2(payload.get("image", ""))
    if img is None:
        return jsonify({"found": False})
    pts = detect_document_contour(img)
    if pts is None:
        return jsonify({"found": False})
    h, w = img.shape[:2]
    return jsonify({"found": True, "points": (pts / [w, h]).tolist()})

@app.route("/scan", methods=["POST"])
def scan():
    payload   = request.get_json(silent=True) or {}
    data      = payload.get("image", "")
    scan_mode = payload.get("mode", "scan")
    img = _b64_to_cv2(data)
    if img is None:
        return jsonify({"ok": False, "error": "Bad image data"})
    pts = detect_document_contour(img)
    if pts is None:
        h, w = img.shape[:2]
        pts = np.array([[0,0],[w,0],[w,h],[0,h]], dtype="float32")
    warped = four_point_transform(img, pts)
    color_bytes, bw_bytes, bw_mat = enhance(warped, mode=scan_mode)
    scanned_bytes = color_bytes if scan_mode == "photo" else bw_bytes
    ocr_text = ""
    warning = None
    try:
        ocr_text = pytesseract.image_to_string(bw_mat, config="--psm 6").strip()
    except TesseractNotFoundError:
        warning = "OCR skipped: Tesseract is not installed or not in PATH."
    except TesseractError as e:
        warning = f"OCR skipped: {str(e)}"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    cv2.imwrite(os.path.join(SAVE_DIR, f"{ts}_color.png"), warped)
    open(os.path.join(SAVE_DIR, f"{ts}_scanned.png"), "wb").write(scanned_bytes)
    return jsonify({"ok": True, "color": _to_b64(color_bytes),
                    "scanned": _to_b64(scanned_bytes), "ocr": ocr_text,
                    "filename": ts, "warning": warning})

@app.route("/outputs/<path:fname>")
def outputs(fname):
    return send_from_directory(SAVE_DIR, fname)

@app.route("/save", methods=["POST"])
def save():
    payload     = request.get_json(silent=True) or {}
    color_b64   = payload.get("color", "")
    scanned_b64 = payload.get("scanned", "")
    ocr_text    = payload.get("ocr", "")
    filename    = payload.get("filename") or datetime.now().strftime("%Y%m%d_%H%M%S")
    if not color_b64 or not scanned_b64:
        return jsonify({"ok": False, "error": "Missing image data"})
    color_bytes   = _b64_to_bytes(color_b64)
    scanned_bytes = _b64_to_bytes(scanned_b64)
    if color_bytes is None or scanned_bytes is None:
        return jsonify({"ok": False, "error": "Bad image data"})
    cp = os.path.join(SAVE_DIR, f"{filename}_color.png")
    sp = os.path.join(SAVE_DIR, f"{filename}_scanned.png")
    tp = os.path.join(SAVE_DIR, f"{filename}_ocr.txt")
    open(cp, "wb").write(color_bytes)
    open(sp, "wb").write(scanned_bytes)
    open(tp, "w", encoding="utf-8").write(ocr_text or "")
    return jsonify({"ok": True,
                    "color":   f"/outputs/{os.path.basename(cp)}",
                    "scanned": f"/outputs/{os.path.basename(sp)}",
                    "ocr":     f"/outputs/{os.path.basename(tp)}"})

# ── Export: plain text ────────────────────────────────────────────────────────
@app.route("/export-txt", methods=["POST"])
def export_txt():
    payload  = request.get_json(silent=True) or {}
    ocr_text = payload.get("ocr", "").strip()
    filename = payload.get("filename") or datetime.now().strftime("%Y%m%d_%H%M%S")
    if not ocr_text:
        return jsonify({"ok": False, "error": "No OCR text to export"}), 400
    ts_label  = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    full_text = f"DocScan Export\nScanned: {ts_label}\n{'='*40}\n\n{ocr_text}\n"
    buf = io.BytesIO(full_text.encode("utf-8"))
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"scan_{filename}.txt",
                     mimetype="text/plain; charset=utf-8")

# ── Export: Word document (.docx) ─────────────────────────────────────────────
@app.route("/export-doc", methods=["POST"])
def export_doc():
    payload   = request.get_json(silent=True) or {}
    ocr_text  = payload.get("ocr", "").strip()
    color_b64 = payload.get("color", "")
    filename  = payload.get("filename") or datetime.now().strftime("%Y%m%d_%H%M%S")
    ts_label  = datetime.now().strftime("%B %d, %Y  %H:%M")

    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Title
    t = doc.add_paragraph()
    tr = t.add_run("Scanned Document")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Timestamp
    ts = doc.add_paragraph()
    tsr = ts.add_run(ts_label)
    tsr.font.size = Pt(9)
    tsr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    ts.paragraph_format.space_after = Pt(14)

    # Embedded image
    if color_b64:
        img_bytes = _b64_to_bytes(color_b64)
        if img_bytes:
            ip = doc.add_paragraph()
            ip.paragraph_format.space_after = Pt(14)
            try:
                ip.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
            except Exception:
                pass

    # Section heading with bottom border
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    sh = doc.add_paragraph()
    shr = sh.add_run("Extracted Text")
    shr.bold = True
    shr.font.size = Pt(12)
    shr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    sh.paragraph_format.space_before = Pt(10)
    sh.paragraph_format.space_after  = Pt(6)
    pPr  = sh._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "4");   bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot); pPr.append(pBdr)

    # OCR lines
    if ocr_text:
        for line in ocr_text.splitlines():
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)
    else:
        p = doc.add_paragraph()
        r = p.add_run("(No text was extracted)")
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"scan_{filename}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ── Helpers ───────────────────────────────────────────────────────────────────
def _b64_to_cv2(b64str):
    try:
        arr = np.frombuffer(base64.b64decode(b64str.split(",")[-1]), np.uint8)
        return cv2.imdecode(arr, cv2.IMREAD_COLOR)
    except Exception:
        return None

def _b64_to_bytes(b64str):
    try:
        return base64.b64decode(b64str.split(",")[-1])
    except Exception:
        return None

def _to_b64(raw_bytes):
    return "data:image/png;base64," + base64.b64encode(raw_bytes).decode()

@app.errorhandler(Exception)
def handle_exception(err):
    if request.path in API_ROUTES:
        app.logger.exception("API error on %s", request.path)
        return jsonify({"ok": False, "error": str(err)}), 500
    if isinstance(err, HTTPException):
        return err
    return "Internal Server Error", 500

if __name__ == "__main__":
    app.run(debug=True, port=5000)