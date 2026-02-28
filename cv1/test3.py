import cv2
import numpy as np
import pytesseract
from PIL import Image
from docx import Document
import os
import datetime

# ----------------------------
# Set Tesseract path (Windows)
# ----------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ----------------------------
# Camera setup (IP Webcam)
# ----------------------------
cap = cv2.VideoCapture(2)
cap.set(3, 1280)
cap.set(4, 720)

# ----------------------------
# Create output directories
# ----------------------------
os.makedirs("scans/images", exist_ok=True)
os.makedirs("scans/docs", exist_ok=True)

# ----------------------------
# Trackbar Setup
# ----------------------------
def empty(a):
    pass

cv2.namedWindow("Parameters")
cv2.resizeWindow("Parameters", 400, 120)
cv2.createTrackbar("Min Area", "Parameters", 13000, 50000, empty)

# ----------------------------
# Utility Functions
# ----------------------------
def reorder(points):
    points = points.reshape((4, 2))
    newPoints = np.zeros((4, 1, 2), np.int32)

    add = points.sum(1)
    newPoints[0] = points[np.argmin(add)]
    newPoints[3] = points[np.argmax(add)]

    diff = np.diff(points, axis=1)
    newPoints[1] = points[np.argmin(diff)]
    newPoints[2] = points[np.argmax(diff)]

    return newPoints


def getBiggestContour(img, minArea):
    biggest = np.array([])
    maxArea = 0

    contours, _ = cv2.findContours(img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    for cnt in contours:
        area = cv2.contourArea(cnt)

        if area > minArea:
            peri = cv2.arcLength(cnt, True)
            approx = cv2.approxPolyDP(cnt, 0.02 * peri, True)

            if len(approx) == 4 and area > maxArea:
                biggest = approx
                maxArea = area

    return biggest


# ----------------------------
# Warp Image + Fix Rotation
# ----------------------------
def warpImageDynamic(img, points):
    points = reorder(points)
    pts = points.reshape(4, 2)

    widthA = np.linalg.norm(pts[0] - pts[1])
    widthB = np.linalg.norm(pts[2] - pts[3])
    width = int(max(widthA, widthB))

    heightA = np.linalg.norm(pts[0] - pts[2])
    heightB = np.linalg.norm(pts[1] - pts[3])
    height = int(max(heightA, heightB))

    dst = np.float32([
        [0, 0],
        [width - 1, 0],
        [0, height - 1],
        [width - 1, height - 1]
    ])

    matrix = cv2.getPerspectiveTransform(np.float32(points), dst)
    warped = cv2.warpPerspective(img, matrix, (width, height))

    # ✅ Fix Horizontal Scan Problem
    if warped.shape[1] > warped.shape[0]:
        warped = cv2.rotate(warped, cv2.ROTATE_90_CLOCKWISE)

    return warped


# ----------------------------
# Enhance Scan Quality
# ----------------------------
def enhanceColorSharp(img):
    lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)

    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    l = clahe.apply(l)

    lab = cv2.merge((l, a, b))
    enhanced = cv2.cvtColor(lab, cv2.COLOR_LAB2BGR)

    kernel = np.array([[0, -1, 0],
                       [-1, 6, -1],
                       [0, -1, 0]])

    enhanced = cv2.filter2D(enhanced, -1, kernel)
    return enhanced


# ----------------------------
# OCR Functions
# ----------------------------
def preprocess_for_ocr(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = cv2.convertScaleAbs(gray, alpha=1.5)

    _, thresh = cv2.threshold(
        gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
    )

    return thresh


def perform_ocr(image):
    processed = preprocess_for_ocr(image)
    pil_img = Image.fromarray(processed)

    config = r"--oem 3 --psm 6 -l eng"
    text = pytesseract.image_to_string(pil_img, config=config)

    return text


# ----------------------------
# Save Image
# ----------------------------
def save_scanned_image(image):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    image_path = f"scans/images/scan_{timestamp}.jpg"

    cv2.imwrite(image_path, image)
    print("✅ Image saved:", image_path)

    return image_path


# ----------------------------
# Save Word Document with Image + Text
# ----------------------------
def save_as_docx(text, image_path):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    doc_path = f"scans/docs/document_{timestamp}.docx"

    doc = Document()
    doc.add_heading("Scanned Document", level=1)

    doc.add_paragraph("Extracted Text:\n")
    doc.add_paragraph(text)

    doc.add_paragraph("\nScanned Image:\n")
    doc.add_picture(image_path, width=None)

    doc.save(doc_path)

    print("✅ DOCX saved:", doc_path)


# ----------------------------
# Main Loop
# ----------------------------
documentCaptured = False
capturedScan = None
scanned_text = ""

print("\n📷 Scanner Ready!")
print("Press 's' to SAVE")
print("Press 'r' to Scan Again")
print("Press 'q' to Quit\n")

while True:
    success, img = cap.read()
    if not success:
        print("❌ Camera not working!")
        break

    areaMin = cv2.getTrackbarPos("Min Area", "Parameters")

    imgBlur = cv2.GaussianBlur(img, (5, 5), 1)
    imgGray = cv2.cvtColor(imgBlur, cv2.COLOR_BGR2GRAY)
    imgCanny = cv2.Canny(imgGray, 50, 150)

    kernel = np.ones((5, 5), np.uint8)
    imgDil = cv2.dilate(imgCanny, kernel, iterations=2)
    imgDil = cv2.erode(imgDil, kernel, iterations=1)

    cv2.imshow("Camera Feed", img)
    cv2.imshow("Edges", imgDil)

    # ----------------------------
    # Capture Document (Preview Only)
    # ----------------------------
    if not documentCaptured:
        biggest = getBiggestContour(imgDil, areaMin)

        if biggest.size != 0:
            imgWarp = warpImageDynamic(img, biggest)
            capturedScan = enhanceColorSharp(imgWarp)

            scanned_text = perform_ocr(capturedScan)

            print("\n📌 Document Captured!")
            print("➡ Press 's' to SAVE")
            print("➡ Press 'r' to scan again\n")

            documentCaptured = True

    # ----------------------------
    # Show Captured Document Preview
    # ----------------------------
    if capturedScan is not None:
        display_img = capturedScan.copy()

        cv2.putText(display_img, "DOCUMENT READY",
                    (20, 50),
                    cv2.FONT_HERSHEY_SIMPLEX,
                    1.2, (0, 255, 0), 3)

        cv2.putText(display_img, "Press 's' to SAVE",
                    (20, 100),
                    cv2.FONT_HERSHEY_SIMPLEX,
                    0.8, (0, 255, 0), 2)

        # OCR Preview Overlay
        lines = scanned_text.split("\n")[:3]
        y0 = 160
        for i, line in enumerate(lines):
            if line.strip():
                cv2.putText(display_img,
                            line[:40],
                            (20, y0 + i * 30),
                            cv2.FONT_HERSHEY_SIMPLEX,
                            0.7, (255, 0, 0), 2)

        cv2.imshow("Scanned Document", display_img)

    # ----------------------------
    # Key Controls
    # ----------------------------
    key = cv2.waitKey(1) & 0xFF

    # Quit
    if key == ord("q"):
        break

    # Reset Scan
    elif key == ord("r"):
        documentCaptured = False
        capturedScan = None
        scanned_text = ""
        print("🔄 Ready for new scan!")

    # Save Manually
    elif key == ord("s") and capturedScan is not None:
        print("💾 Saving document...")

        img_path = save_scanned_image(capturedScan)
        save_as_docx(scanned_text, img_path)

        print("✅ Saved Successfully!\n")


# ----------------------------
# Cleanup
# ----------------------------
cap.release()
cv2.destroyAllWindows()
